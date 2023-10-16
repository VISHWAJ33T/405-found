import re
import aspose.words as aw
import docx
# doc = aw.Document("./Samples/sample_level_1.html")
# doc = aw.Document("./Samples/sample_level_2.html")
# doc = aw.Document("./Samples/sample_level_3.html")
doc = aw.Document("./Samples/sample_level_4.html")


# doc.save("Output/without watermark/sample_level_1.docx")
# doc.save("Output/without watermark/sample_level_2.docx")
# doc.save("Output/without watermark/sample_level_3.docx")
doc.save("Output/without watermark/sample_level_4.docx")


print("Word files without Watermark saved successfully")


# document = docx.Document("Output/without watermark/sample_level_1.docx")
# document = docx.Document("Output/without watermark/sample_level_2.docx")
# document = docx.Document("Output/without watermark/sample_level_3.docx")
document = docx.Document("Output/without watermark/sample_level_4.docx")

waterMark = document.paragraphs[1].text

# Instantiate and fill the TextWatermarkOptions object
options = aw.TextWatermarkOptions()
options.font_family = "Calibri"
options.font_size = 42
options.layout = aw.WatermarkLayout.DIAGONAL
options.is_semitrasparent = True

doc.watermark.set_text(waterMark, options)
# Save the document
# doc.save("Output/sample_level_1.docx")
# doc.save("Output/sample_level_2.docx")
# doc.save("Output/sample_level_3.docx")
doc.save("Output/sample_level_4.docx")

print("Watermark added successfully in the Word file")
